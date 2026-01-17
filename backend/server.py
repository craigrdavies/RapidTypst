from fastapi import FastAPI, APIRouter, HTTPException
from fastapi.responses import FileResponse, Response
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional
import uuid
from datetime import datetime, timezone
import tempfile
import typst
import subprocess
import shutil

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Temp directory for typst files
TEMP_DIR = Path(tempfile.gettempdir()) / "typst_editor"
TEMP_DIR.mkdir(exist_ok=True)


# Define Models
class DocumentCreate(BaseModel):
    title: str
    content: str


class DocumentUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None


class Document(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    content: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class CompileRequest(BaseModel):
    content: str


class CompileResponse(BaseModel):
    success: bool
    html: Optional[str] = None
    error: Optional[str] = None


class ExportRequest(BaseModel):
    content: str
    format: str  # 'pdf', 'html', 'docx'


# Helper function to compile typst
def compile_typst_to_pdf(content: str, output_path: Path) -> tuple[bool, Optional[str]]:
    """Compile typst content to PDF using the typst Python package"""
    try:
        # Write content to a temp .typ file
        typ_file = output_path.with_suffix('.typ')
        typ_file.write_text(content, encoding='utf-8')
        
        # Compile using typst Python package
        pdf_bytes = typst.compile(str(typ_file))
        output_path.write_bytes(pdf_bytes)
        
        return True, None
    except Exception as e:
        return False, str(e)


def compile_typst_to_svg(content: str) -> tuple[bool, Optional[str], Optional[str]]:
    """Compile typst content to SVG for preview"""
    try:
        # Create temp files
        typ_file = TEMP_DIR / f"{uuid.uuid4()}.typ"
        svg_dir = TEMP_DIR / f"svg_{uuid.uuid4()}"
        svg_dir.mkdir(exist_ok=True)
        
        typ_file.write_text(content, encoding='utf-8')
        
        # Try using typst CLI for SVG output
        result = subprocess.run(
            ['typst', 'compile', str(typ_file), str(svg_dir / 'page{n}.svg')],
            capture_output=True,
            text=True,
            timeout=30
        )
        
        if result.returncode != 0:
            # Clean up
            typ_file.unlink(missing_ok=True)
            shutil.rmtree(svg_dir, ignore_errors=True)
            return False, None, result.stderr or "Compilation failed"
        
        # Read all SVG files and combine them
        svg_files = sorted(svg_dir.glob('*.svg'))
        svgs = []
        for svg_file in svg_files:
            svg_content = svg_file.read_text(encoding='utf-8')
            svgs.append(svg_content)
        
        # Clean up
        typ_file.unlink(missing_ok=True)
        shutil.rmtree(svg_dir, ignore_errors=True)
        
        if svgs:
            # Wrap SVGs in HTML
            html_content = f"""
            <div style="display: flex; flex-direction: column; gap: 20px; padding: 20px; background: white;">
                {''.join(f'<div style="box-shadow: 0 2px 8px rgba(0,0,0,0.1); padding: 10px; background: white;">{svg}</div>' for svg in svgs)}
            </div>
            """
            return True, html_content, None
        else:
            return False, None, "No output generated"
            
    except subprocess.TimeoutExpired:
        return False, None, "Compilation timed out"
    except FileNotFoundError:
        # Typst CLI not installed, fallback to PDF preview message
        return False, None, "Typst CLI not found. Install it for live preview."
    except Exception as e:
        return False, None, str(e)


# API Routes
@api_router.get("/")
async def root():
    return {"message": "Rapid Typst API"}


# Document CRUD
@api_router.post("/documents", response_model=Document)
async def create_document(doc: DocumentCreate):
    document = Document(title=doc.title, content=doc.content)
    doc_dict = document.model_dump()
    doc_dict['created_at'] = doc_dict['created_at'].isoformat()
    doc_dict['updated_at'] = doc_dict['updated_at'].isoformat()
    
    await db.documents.insert_one(doc_dict)
    return document


@api_router.get("/documents", response_model=List[Document])
async def list_documents():
    docs = await db.documents.find({}, {"_id": 0}).to_list(1000)
    for doc in docs:
        if isinstance(doc.get('created_at'), str):
            doc['created_at'] = datetime.fromisoformat(doc['created_at'])
        if isinstance(doc.get('updated_at'), str):
            doc['updated_at'] = datetime.fromisoformat(doc['updated_at'])
    return docs


@api_router.get("/documents/{doc_id}", response_model=Document)
async def get_document(doc_id: str):
    doc = await db.documents.find_one({"id": doc_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if isinstance(doc.get('created_at'), str):
        doc['created_at'] = datetime.fromisoformat(doc['created_at'])
    if isinstance(doc.get('updated_at'), str):
        doc['updated_at'] = datetime.fromisoformat(doc['updated_at'])
    return doc


@api_router.put("/documents/{doc_id}", response_model=Document)
async def update_document(doc_id: str, update: DocumentUpdate):
    doc = await db.documents.find_one({"id": doc_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    update_data = {k: v for k, v in update.model_dump().items() if v is not None}
    update_data['updated_at'] = datetime.now(timezone.utc).isoformat()
    
    await db.documents.update_one({"id": doc_id}, {"$set": update_data})
    
    updated_doc = await db.documents.find_one({"id": doc_id}, {"_id": 0})
    if isinstance(updated_doc.get('created_at'), str):
        updated_doc['created_at'] = datetime.fromisoformat(updated_doc['created_at'])
    if isinstance(updated_doc.get('updated_at'), str):
        updated_doc['updated_at'] = datetime.fromisoformat(updated_doc['updated_at'])
    return updated_doc


@api_router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    result = await db.documents.delete_one({"id": doc_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"message": "Document deleted"}


# Compile endpoint for live preview
@api_router.post("/compile", response_model=CompileResponse)
async def compile_typst(request: CompileRequest):
    if not request.content.strip():
        return CompileResponse(
            success=True, 
            html='<div style="color: #71717A; padding: 40px; text-align: center;">Start typing Typst markup to see preview...</div>'
        )
    
    success, html, error = compile_typst_to_svg(request.content)
    
    if success:
        return CompileResponse(success=True, html=html)
    else:
        # Return a styled error message
        error_html = f'''
        <div style="padding: 20px; background: #FEF2F2; border: 1px solid #FECACA; border-radius: 4px; margin: 20px;">
            <div style="color: #DC2626; font-weight: 600; margin-bottom: 8px;">Compilation Error</div>
            <pre style="color: #991B1B; font-size: 13px; white-space: pre-wrap; margin: 0; font-family: 'JetBrains Mono', monospace;">{error}</pre>
        </div>
        '''
        return CompileResponse(success=False, html=error_html, error=error)


# Export endpoints
@api_router.post("/export/pdf")
async def export_pdf(request: ExportRequest):
    try:
        output_path = TEMP_DIR / f"{uuid.uuid4()}.pdf"
        success, error = compile_typst_to_pdf(request.content, output_path)
        
        if not success:
            raise HTTPException(status_code=400, detail=error)
        
        return FileResponse(
            path=str(output_path),
            media_type='application/pdf',
            filename='document.pdf'
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post("/export/html")
async def export_html(request: ExportRequest):
    try:
        success, html, error = compile_typst_to_svg(request.content)
        
        if not success:
            raise HTTPException(status_code=400, detail=error)
        
        # Create a full HTML document
        full_html = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Typst Document</title>
    <style>
        body {{ margin: 0; padding: 20px; background: #f5f5f5; font-family: system-ui, sans-serif; }}
        .page {{ background: white; box-shadow: 0 2px 8px rgba(0,0,0,0.1); margin: 0 auto 20px; max-width: 800px; }}
    </style>
</head>
<body>
{html}
</body>
</html>'''
        
        return Response(
            content=full_html,
            media_type='text/html',
            headers={'Content-Disposition': 'attachment; filename="document.html"'}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post("/export/docx")
async def export_docx(request: ExportRequest):
    try:
        # First compile to PDF, then try to convert
        from docx import Document as DocxDocument
        from docx.shared import Pt
        
        # Create a simple docx with the typst content as text
        # Note: Full typst->docx conversion is complex; this is a basic version
        doc = DocxDocument()
        
        # Add a note about the conversion
        para = doc.add_paragraph()
        run = para.add_run("Note: This is a basic export. For best results, use PDF export.")
        run.italic = True
        run.font.size = Pt(10)
        
        doc.add_paragraph()  # Spacer
        
        # Add the typst source
        para = doc.add_paragraph()
        run = para.add_run("Typst Source:")
        run.bold = True
        
        para = doc.add_paragraph()
        para.add_run(request.content)
        
        # Save to temp file
        output_path = TEMP_DIR / f"{uuid.uuid4()}.docx"
        doc.save(str(output_path))
        
        return FileResponse(
            path=str(output_path),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename='document.docx'
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
